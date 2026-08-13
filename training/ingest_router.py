#!/usr/bin/env python3
"""Ingest router — Stage 0 of the curation pipeline.

Dispatches incoming source records to type-specific extractors (web, document,
api, youtube) and streams raw JSONL shards to ``data/raw/<source_type>/``.

Each emitted record carries::

    {
      "id": "<unique-id>",
      "source_type": "web|document|api|youtube",
      "source_url": "<url or file path>",
      "raw_text": "<extracted text content>",
      "metadata": {
        "license": "<SPDX id>",
        "fetch_ts": "<ISO-8601 UTC>",
        "lang": "<language hint>",
        ...
      },
      "provenance": { ... }   # validated via training.provenance
    }

Shards are capped at ``SHARD_SIZE`` records (50 000 by default) and named
``shard-NNNNN.jsonl``.

The router reads a source manifest (``data/licenses/source_manifest.yaml``)
at startup to populate its license gate.  Records whose license does not
validate against ``training.provenance.ALLOWED_LICENSES`` are dropped with a
logged reason.
"""

from __future__ import annotations

import asyncio
import hashlib
import json
import logging
import os
import re
import subprocess
import time
from collections.abc import AsyncIterator, Awaitable, Callable, Mapping, Sequence
from dataclasses import dataclass, field
from datetime import UTC, datetime
from pathlib import Path
from typing import IO, Any
from urllib.parse import urlparse

import httpx
import yaml
from selectolax.parser import HTMLParser

from training.provenance import (
    ProvenanceOptions,
    build_provenance,
    validate_license,
)

logger = logging.getLogger("ingest_router")

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

SHARD_SIZE = 50_000
ROBOTS_CACHE_TTL_SECONDS = 24 * 3600  # 24 hours
WEB_RATE_LIMIT_SECONDS = 2.0  # 1 request per 2 seconds per domain

# Reuse NEMO_RETRY_DELAYS from the repo if available; otherwise use the default
# specified in the task brief.
try:
    from training.sdg_pipeline import NEMO_RETRY_DELAYS  # type: ignore[import-not-found]
except Exception:  # noqa: BLE001  (import fallback is best-effort)
    NEMO_RETRY_DELAYS: tuple[float, ...] = (0.5, 1.0, 2.0, 4.0)

DEFAULT_MANIFEST_PATH = "data/licenses/source_manifest.yaml"
DEFAULT_RAW_OUTPUT_DIR = Path("data/raw")


# ---------------------------------------------------------------------------
# Source manifest
# ---------------------------------------------------------------------------


@dataclass
class ManifestEntry:
    """A single source-to-license mapping from the source manifest."""

    url_or_domain: str
    license_id: str
    source_type: str = "web"
    provenance: dict[str, Any] = field(default_factory=dict)

    def matches(self, url: str) -> bool:
        """Return True if *url* falls under this entry (domain or prefix match)."""
        if self.url_or_domain in url:
            return True
        try:
            parsed = urlparse(url)
            netloc = parsed.netloc.lower()
            return self.url_or_domain.lower() in netloc
        except Exception:  # noqa: BLE001
            return False


def load_source_manifest(manifest_path: Path | str) -> dict[str, ManifestEntry]:
    """Load the YAML source manifest into a dict keyed by url_or_domain.

    Returns an empty dict if the file is missing (the license gate will then
    default to ``NOASSERTION`` with a warning).
    """

    path = Path(manifest_path)
    if not path.exists():
        logger.warning("Source manifest not found at %s — license gate defaults to NOASSERTION", path)
        return {}
    data = yaml.safe_load(path.read_text(encoding="utf-8")) or {}
    entries: dict[str, ManifestEntry] = {}
    for item in data.get("sources", []):
        key = item["url"]
        entries[key] = ManifestEntry(
            url_or_domain=key,
            license_id=item.get("license", "NOASSERTION"),
            source_type=item.get("source_type", "web"),
            provenance=item.get("provenance", {}),
        )
    return entries


def lookup_license(
    source_url: str,
    manifest: Mapping[str, ManifestEntry],
) -> str:
    """Look up the license for *source_url* in the manifest.

    Falls back to ``NOASSERTION`` if no entry matches.
    """

    for entry in manifest.values():
        if entry.matches(source_url):
            return entry.license_id
    return "NOASSERTION"


# ---------------------------------------------------------------------------
# Shard writer
# ---------------------------------------------------------------------------


class ShardWriter:
    """Stream JSONL records into shard files capped at ``SHARD_SIZE`` records."""

    def __init__(self, output_dir: Path, source_type: str, shard_size: int = SHARD_SIZE) -> None:
        self._output_dir = output_dir
        self._source_type = source_type
        self._shard_size = shard_size
        self._shard_index = 0
        self._records_in_shard = 0
        self._fh: IO[str] | None = None
        self._total_written = 0

    def _shard_path(self) -> Path:
        return self._output_dir / f"shard-{self._shard_index:05d}.jsonl"

    def _open_new_shard(self) -> None:
        self._output_dir.mkdir(parents=True, exist_ok=True)
        if self._fh is not None:
            self._fh.close()
        path = self._shard_path()
        self._fh = open(path, "w", encoding="utf-8")
        self._records_in_shard = 0
        logger.info("Opened shard %s for source_type=%s", path, self._source_type)

    def write(self, record: Mapping[str, Any]) -> None:
        if self._fh is None or self._records_in_shard >= self._shard_size:
            self._shard_index += 1
            self._open_new_shard()
        assert self._fh is not None  # _open_new_shard guarantees non-None
        self._fh.write(json.dumps(record, ensure_ascii=False) + "\n")
        self._records_in_shard += 1
        self._total_written += 1

    def close(self) -> None:
        if self._fh is not None:
            self._fh.close()
            self._fh = None
        logger.info(
            "Shard writer closed: %d records written across %d shards for source_type=%s",
            self._total_written,
            self._shard_index,
            self._source_type,
        )

    @property
    def total_written(self) -> int:
        return self._total_written

    @property
    def shard_count(self) -> int:
        return self._shard_index


# ---------------------------------------------------------------------------
# License gate
# ---------------------------------------------------------------------------


class LicenseGate:
    """Validates record licenses before emission.  Drops invalid records."""

    def __init__(self, manifest: Mapping[str, ManifestEntry]) -> None:
        self._manifest = manifest
        self._dropped = 0

    def check(self, source_url: str) -> tuple[bool, str]:
        """Return ``(ok, license_id)``.  When *ok* is False the record is dropped."""
        license_id = lookup_license(source_url, self._manifest)
        try:
            validated = validate_license(license_id)
            return True, validated
        except ValueError as exc:
            self._dropped += 1
            logger.warning("License gate dropped record from %s: %s", source_url, exc)
            return False, license_id

    @property
    def dropped_count(self) -> int:
        return self._dropped


# ---------------------------------------------------------------------------
# Record builders
# ---------------------------------------------------------------------------


def _utc_now_iso() -> str:
    return datetime.now(UTC).isoformat()


def _record_id(source_url: str, raw_text: str) -> str:
    """Deterministic record ID from source URL + content hash."""
    digest = hashlib.sha256(f"{source_url}:{raw_text[:512]}".encode("utf-8")).hexdigest()
    return digest[:16]


def build_record(
    *,
    source_type: str,
    source_url: str,
    raw_text: str,
    metadata: Mapping[str, Any] | None = None,
    license_id: str = "NOASSERTION",
) -> dict[str, Any]:
    """Build a fully-formed record with provenance for shard emission."""

    record_id = _record_id(source_url, raw_text)
    meta: dict[str, Any] = {
        "license": license_id,
        "fetch_ts": _utc_now_iso(),
        "lang": _detect_lang_hint(raw_text),
    }
    if metadata:
        meta.update(metadata)

    provenance = build_provenance(
        source_url,
        source_type,
        options=ProvenanceOptions(license_id=license_id),
        metadata=meta,
    )

    return {
        "id": record_id,
        "source_type": source_type,
        "source_url": source_url,
        "raw_text": raw_text,
        "metadata": meta,
        "provenance": provenance,
    }


def _detect_lang_hint(text: str) -> str:
    """Cheap heuristic: detect German vs English based on common words."""
    if not text:
        return "unknown"
    sample = text[:2000].lower()
    german_markers = (" der ", " die ", " das ", " und ", " nicht ", " ist ", " ein ", " eine ")
    english_markers = (" the ", " and ", " is ", " not ", " a ", " to ", " of ")
    german_count = sum(sample.count(m) for m in german_markers)
    english_count = sum(sample.count(m) for m in english_markers)
    if german_count > english_count:
        return "de"
    return "en"


# ---------------------------------------------------------------------------
# Web extractor
# ---------------------------------------------------------------------------


class WebExtractor:
    """Fetches web pages, extracts main content via trafilatura.

    Features:
    - ``httpx.AsyncClient`` with connection pooling.
    - robots.txt caching per domain (24h TTL).
    - Per-domain rate limiting (1 request per 2 seconds).
    - ``selectolax`` for HTML parsing, ``trafilatura`` for main-content extraction.
    """

    def __init__(
        self,
        *,
        rate_limit_seconds: float = WEB_RATE_LIMIT_SECONDS,
        robots_ttl: int = ROBOTS_CACHE_TTL_SECONDS,
        client: httpx.AsyncClient | None = None,
    ) -> None:
        self._rate_limit = rate_limit_seconds
        self._robots_ttl = robots_ttl
        self._client = client
        self._robots_cache: dict[str, tuple[float, bool]] = {}
        self._domain_last_request: dict[str, float] = {}
        self._domain_locks: dict[str, asyncio.Lock] = {}

    async def _get_client(self) -> httpx.AsyncClient:
        if self._client is None or self._client.is_closed:
            self._client = httpx.AsyncClient(
                timeout=httpx.Timeout(30.0, connect=10.0),
                limits=httpx.Limits(max_connections=20, max_keepalive_connections=10),
                follow_redirects=True,
            )
        return self._client

    def _domain(self, url: str) -> str:
        return urlparse(url).netloc.lower()

    def _get_domain_lock(self, domain: str) -> asyncio.Lock:
        if domain not in self._domain_locks:
            self._domain_locks[domain] = asyncio.Lock()
        return self._domain_locks[domain]

    async def _check_robots(self, url: str, client: httpx.AsyncClient) -> bool:
        """Return True if *url* is allowed by robots.txt (cached, 24h TTL)."""
        domain = self._domain(url)
        now = time.monotonic()
        cached = self._robots_cache.get(domain)
        if cached is not None:
            ts, allowed = cached
            if now - ts < self._robots_ttl:
                return allowed

        robots_url = f"https://{domain}/robots.txt"
        allowed = True  # permissive default
        try:
            resp = await client.get(robots_url)
            if resp.status_code == 200:
                # Simple check: if Disallow: / is present, block all.
                # This is a simplified robots.txt parser.
                text = resp.text
                pattern = r"User-agent:\s*\*.*?(?:(?:User-agent:)|\Z)"
                for block in re.finditer(pattern, text, re.IGNORECASE | re.DOTALL):
                    block_text = block.group()
                    if re.search(r"Disallow:\s*/\s*$", block_text, re.IGNORECASE | re.MULTILINE):
                        allowed = False
                        break
        except httpx.HTTPError:
            pass  # be permissive on error

        self._robots_cache[domain] = (now, allowed)
        return allowed

    async def _enforce_rate_limit(self, domain: str) -> None:
        lock = self._get_domain_lock(domain)
        async with lock:
            last = self._domain_last_request.get(domain, 0.0)
            elapsed = time.monotonic() - last
            if elapsed < self._rate_limit:
                await asyncio.sleep(self._rate_limit - elapsed)
            self._domain_last_request[domain] = time.monotonic()

    async def extract(self, source_url: str) -> dict[str, Any]:
        """Fetch *source_url*, extract main content. Returns record fields."""
        client = await self._get_client()
        domain = self._domain(source_url)

        # robots.txt check
        if not await self._check_robots(source_url, client):
            logger.info("robots.txt disallows %s — skipping", source_url)
            return {
                "source_url": source_url,
                "raw_text": "",
                "metadata": {"fetch_status": 403, "html_title": ""},
            }

        await self._enforce_rate_limit(domain)

        try:
            resp = await client.get(source_url, headers={"User-Agent": "PixelatedAI-IngestRouter/1.0"})
            status = resp.status_code
            html = resp.text
        except httpx.HTTPError as exc:
            logger.warning("HTTP error fetching %s: %s", source_url, exc)
            return {
                "source_url": source_url,
                "raw_text": "",
                "metadata": {"fetch_status": 0, "html_title": ""},
            }

        # Extract title with selectolax
        tree = HTMLParser(html)
        title_node = tree.css_first("title")
        html_title = title_node.text(strip=True) if title_node else ""

        # Extract main content with trafilatura
        raw_text = self._trafilatura_extract(html)

        return {
            "source_url": source_url,
            "raw_text": raw_text,
            "metadata": {
                "fetch_status": status,
                "html_title": html_title,
            },
        }

    @staticmethod
    def _trafilatura_extract(html: str) -> str:
        """Extract main content from HTML using trafilatura."""
        try:
            import trafilatura

            return trafilatura.extract(html) or ""
        except ImportError:
            logger.warning("trafilatura not installed — falling back to selectolax text extraction")
            tree = HTMLParser(html)
            return tree.body.text(separator="\n", strip=True) if tree.body else ""

    async def close(self) -> None:
        if self._client is not None and not self._client.is_closed:
            await self._client.aclose()


# ---------------------------------------------------------------------------
# Document extractor
# ---------------------------------------------------------------------------


class DocumentExtractor:
    """Extracts text from PDF, DOCX, and HTML files.

    - PDF: delegates to the existing ``book_pdf_converter.py`` pattern.
    - DOCX: uses ``python-docx`` to extract paragraphs + tables as markdown.
    - HTML: reuses the web extractor's trafilatura path.
    """

    def __init__(self, web_extractor: WebExtractor | None = None) -> None:
        self._web_extractor = web_extractor or WebExtractor()

    async def extract(self, source_ref: str) -> dict[str, Any]:
        """Extract text from a local file path (*source_ref*)."""
        path = Path(source_ref)
        suffix = path.suffix.lower()

        if suffix in {".pdf"}:
            return self._extract_pdf(path)
        if suffix in {".docx"}:
            return self._extract_docx(path)
        if suffix in {".html", ".htm"}:
            return await self._extract_html(path)
        raise ValueError(f"Unsupported document type: {suffix}")

    def _extract_pdf(self, path: Path) -> dict[str, Any]:
        """Extract text from PDF using pypdf (same as book_pdf_converter.py)."""
        try:
            from pypdf import PdfReader
        except ImportError:
            logger.error("pypdf not installed — cannot extract PDF")
            return {
                "source_url": str(path),
                "raw_text": "",
                "metadata": {"format": "pdf", "error": "pypdf unavailable"},
            }

        reader = PdfReader(str(path))
        pages: list[str] = []
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                pages.append(text.strip())
        raw_text = "\n\n".join(pages)
        return {
            "source_url": str(path),
            "raw_text": raw_text,
            "metadata": {"format": "pdf", "page_count": len(reader.pages)},
        }

    def _extract_docx(self, path: Path) -> dict[str, Any]:
        """Extract text from DOCX preserving heading structure as markdown."""
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx not installed — cannot extract DOCX")
            return {
                "source_url": str(path),
                "raw_text": "",
                "metadata": {"format": "docx", "error": "python-docx unavailable"},
            }

        doc = Document(str(path))
        lines: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = (para.style.name or "").lower() if para.style else ""
            if "heading 1" in style_name:
                lines.append(f"# {text}")
            elif "heading 2" in style_name:
                lines.append(f"## {text}")
            elif "heading 3" in style_name:
                lines.append(f"### {text}")
            elif "heading" in style_name:
                # Generic heading — extract level if possible
                level_match = re.search(r"heading\s*(\d+)", style_name)
                level = int(level_match.group(1)) if level_match else 4
                lines.append(f"{'#' * min(level, 6)} {text}")
            else:
                lines.append(text)

        # Extract tables as markdown
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")

        raw_text = "\n\n".join(lines)
        return {
            "source_url": str(path),
            "raw_text": raw_text,
            "metadata": {"format": "docx", "paragraph_count": len(doc.paragraphs), "table_count": len(doc.tables)},
        }

    async def _extract_html(self, path: Path) -> dict[str, Any]:
        """Extract text from a local HTML file using trafilatura."""
        html = path.read_text(encoding="utf-8", errors="replace")
        raw_text = WebExtractor._trafilatura_extract(html)
        return {
            "source_url": str(path),
            "raw_text": raw_text,
            "metadata": {"format": "html"},
        }

    async def close(self) -> None:
        await self._web_extractor.close()


# ---------------------------------------------------------------------------
# API extractor
# ---------------------------------------------------------------------------


@dataclass
class APIExtractorConfig:
    """Configuration for the API extractor."""

    endpoint: str
    headers: dict[str, str] = field(default_factory=dict)
    params: dict[str, str] = field(default_factory=dict)
    text_field: str = "text"
    cursor_field: str = "cursor"
    next_cursor_field: str = "next_cursor"
    items_field: str = "items"
    page_size: int = 100
    max_pages: int = 1000


class APIExtractor:
    """Fetches records from an API with exponential backoff and pagination.

    - ``httpx.AsyncClient`` with ``NEMO_RETRY_DELAYS``-style backoff.
    - Configurable endpoint, headers, pagination (cursor + offset).
    - Emits one record per API item.
    """

    def __init__(
        self,
        config: APIExtractorConfig,
        *,
        retry_delays: Sequence[float] = NEMO_RETRY_DELAYS,
        client: httpx.AsyncClient | None = None,
    ) -> None:
        self._config = config
        self._retry_delays = retry_delays
        self._client = client

    async def _get_client(self) -> httpx.AsyncClient:
        if self._client is None or self._client.is_closed:
            self._client = httpx.AsyncClient(
                timeout=httpx.Timeout(60.0, connect=15.0),
                limits=httpx.Limits(max_connections=10, max_keepalive_connections=5),
                follow_redirects=True,
            )
        return self._client

    async def _fetch_with_retry(self, url: str, params: dict[str, Any]) -> httpx.Response | None:
        client = await self._get_client()
        last_exc: Exception | None = None
        for attempt, delay in enumerate(self._retry_delays):
            try:
                resp = await client.get(url, params=params, headers=self._config.headers)
                if resp.status_code == 429 or resp.status_code >= 500:
                    logger.warning(
                        "API returned %d for %s (attempt %d) — retrying in %.1fs",
                        resp.status_code,
                        url,
                        attempt + 1,
                        delay,
                    )
                    await asyncio.sleep(delay)
                    continue
                resp.raise_for_status()
                return resp
            except httpx.HTTPError as exc:
                last_exc = exc
                logger.warning(
                    "API request failed for %s (attempt %d): %s — retrying in %.1fs", url, attempt + 1, exc, delay
                )
                await asyncio.sleep(delay)
        if last_exc is not None:
            logger.error("API request exhausted retries for %s: %s", url, last_exc)
        return None

    async def extract(self, source_url: str) -> AsyncIterator[dict[str, Any]]:
        """Yield one record dict per API item.

        *source_url* is used as the base URL for provenance; the actual API
        endpoint comes from the config.
        """
        endpoint = self._config.endpoint
        cursor: str | None = None
        page = 0

        while page < self._config.max_pages:
            params = dict(self._config.params)
            params["limit"] = str(self._config.page_size)
            if cursor is not None:
                params[self._config.cursor_field] = cursor

            resp = await self._fetch_with_retry(endpoint, params)
            if resp is None:
                break

            data = resp.json()
            items = data.get(self._config.items_field, [])
            if not items:
                break

            for item in items:
                raw_text = str(item.get(self._config.text_field, ""))
                item_id = item.get("id", _record_id(source_url, raw_text))
                yield {
                    "source_url": source_url,
                    "raw_text": raw_text,
                    "metadata": {
                        "api_item_id": str(item_id),
                        "api_page": page,
                        **{k: v for k, v in item.items() if k != self._config.text_field and k != "id"},
                    },
                    "id_override": str(item_id),
                }

            # Check for next cursor
            next_cursor = data.get(self._config.next_cursor_field)
            if next_cursor is None or next_cursor == cursor:
                break
            cursor = next_cursor
            page += 1

    async def close(self) -> None:
        if self._client is not None and not self._client.is_closed:
            await self._client.aclose()


# ---------------------------------------------------------------------------
# YouTube extractor
# ---------------------------------------------------------------------------


class YouTubeExtractor:
    """Fetches YouTube video transcripts via ``yt-dlp``.

    Delegates to the existing ``transcript_fetcher.py`` pattern where possible;
    otherwise implements a thin ``yt-dlp`` wrapper that fetches transcript text.
    """

    def __init__(self, *, yt_dlp_timeout: int = 60) -> None:
        self._timeout = yt_dlp_timeout

    @staticmethod
    def _extract_video_id(url: str) -> str:
        """Extract YouTube video ID from various URL formats."""
        if "youtu.be/" in url:
            return url.split("youtu.be/")[1].split("?")[0]
        if "watch?v=" in url:
            return url.split("watch?v=")[1].split("&")[0]
        if "/embed/" in url:
            return url.split("/embed/")[1].split("?")[0]
        # Fallback — return as-is
        return url

    async def extract(self, source_url: str) -> dict[str, Any]:
        """Fetch transcript for a single YouTube video URL."""
        video_id = self._extract_video_id(source_url)

        # Run yt-dlp in a thread to avoid blocking the event loop
        loop = asyncio.get_event_loop()
        result = await loop.run_in_executor(
            None,
            self._fetch_transcript_sync,
            video_id,
            source_url,
        )
        return result

    def _fetch_transcript_sync(self, video_id: str, source_url: str) -> dict[str, Any]:
        """Fetch transcript using yt-dlp subprocess (blocking)."""
        cmd = [
            "yt-dlp",
            "--write-auto-sub",
            "--sub-lang",
            "en,en-US,en-GB,de,de-DE",
            "--skip-download",
            "--output",
            "/tmp/yt-transcript-%(id)s",
            "--user-agent",
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
            "--sleep-interval",
            "1",
            "--max-sleep-interval",
            "5",
            f"https://www.youtube.com/watch?v={video_id}",
        ]

        try:
            subprocess.run(cmd, capture_output=True, text=True, timeout=self._timeout, check=False)
        except (subprocess.TimeoutExpired, FileNotFoundError) as exc:
            logger.warning("yt-dlp failed for %s: %s", video_id, exc)
            return {
                "source_url": source_url,
                "raw_text": "",
                "metadata": {"video_id": video_id, "duration": 0, "error": str(exc)},
            }

        # Find the generated subtitle file
        import glob

        patterns = [
            f"/tmp/yt-transcript-{video_id}*.vtt",
            f"/tmp/yt-transcript-{video_id}*.srt",
            f"/tmp/yt-transcript-{video_id}*.srv3",
            f"/tmp/yt-transcript-{video_id}*.srv2",
            f"/tmp/yt-transcript-{video_id}*.srv1",
        ]
        subtitle_path: str | None = None
        for pat in patterns:
            files = glob.glob(pat)
            if files:
                subtitle_path = files[0]
                break

        if subtitle_path is None:
            logger.warning("No transcript found for video %s", video_id)
            return {
                "source_url": source_url,
                "raw_text": "",
                "metadata": {"video_id": video_id, "duration": 0, "error": "no transcript found"},
            }

        # Read and clean subtitle text
        raw_text = self._clean_subtitle(Path(subtitle_path).read_text(encoding="utf-8", errors="replace"))

        # Get duration via yt-dlp --dump-json
        duration = 0
        try:
            result = subprocess.run(
                ["yt-dlp", "--dump-json", "--skip-download", f"https://www.youtube.com/watch?v={video_id}"],
                capture_output=True,
                text=True,
                timeout=self._timeout,
                check=False,
            )
            if result.returncode == 0 and result.stdout.strip():
                info = json.loads(result.stdout.strip().splitlines()[0])
                duration = int(info.get("duration", 0))
        except (subprocess.TimeoutExpired, json.JSONDecodeError, FileNotFoundError):
            pass

        return {
            "source_url": source_url,
            "raw_text": raw_text,
            "metadata": {"video_id": video_id, "duration": duration},
        }

    @staticmethod
    def _clean_subtitle(text: str) -> str:
        """Strip subtitle formatting (timestamps, HTML tags) and return plain text."""
        # Remove VTT/SRT headers and timestamps
        text = re.sub(r"WEBVTT.*?\n\n", "", text, flags=re.DOTALL)
        text = re.sub(r"\d{2}:\d{2}:\d{2}[.,]\d{3}\s*-->\s*\d{2}:\d{2}:\d{2}[.,]\d{3}", "", text)
        text = re.sub(r"<[^>]+>", "", text)  # HTML tags
        text = re.sub(r"^\d+\s*$", "", text, flags=re.MULTILINE)  # subtitle indices
        # Collapse whitespace
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        return " ".join(lines)

    async def close(self) -> None:
        pass  # no persistent resources


# ---------------------------------------------------------------------------
# Ingest router
# ---------------------------------------------------------------------------


# Type alias for async extractor functions that return a single record dict
SingleRecordExtractor = Callable[[str], Awaitable[dict[str, Any]]]
# Type alias for async extractor functions that yield multiple record dicts
MultiRecordExtractor = Callable[[str], AsyncIterator[dict[str, Any]]]


class IngestRouter:
    """Dispatches incoming source records to type-specific extractors.

    Usage::

        router = IngestRouter(manifest_path="data/licenses/source_manifest.yaml")
        await router.ingest([
            {"source_type": "web", "source_url": "https://example.com/page"},
            {"source_type": "document", "source_url": "/path/to/file.pdf"},
            {"source_type": "api", "source_url": "https://api.example.com/data"},
            {"source_type": "youtube", "source_url": "https://youtube.com/watch?v=..."},
        ])

    Records are streamed to ``data/raw/<source_type>/shard-NNNNN.jsonl`` (50K per shard).
    """

    SOURCE_TYPES: frozenset[str] = frozenset({"web", "document", "api", "youtube"})

    def __init__(
        self,
        *,
        manifest_path: Path | str = DEFAULT_MANIFEST_PATH,
        raw_output_dir: Path = DEFAULT_RAW_OUTPUT_DIR,
        shard_size: int = SHARD_SIZE,
        web_extractor: WebExtractor | None = None,
        document_extractor: DocumentExtractor | None = None,
        api_extractor: APIExtractor | None = None,
        youtube_extractor: YouTubeExtractor | None = None,
    ) -> None:
        self._manifest = load_source_manifest(manifest_path)
        self._license_gate = LicenseGate(self._manifest)
        self._raw_output_dir = raw_output_dir
        self._shard_size = shard_size

        self._web_extractor = web_extractor or WebExtractor()
        self._document_extractor = document_extractor or DocumentExtractor(self._web_extractor)
        self._api_extractor = api_extractor
        self._youtube_extractor = youtube_extractor or YouTubeExtractor()

        self._shard_writers: dict[str, ShardWriter] = {}

    def _get_shard_writer(self, source_type: str) -> ShardWriter:
        if source_type not in self._shard_writers:
            output_dir = self._raw_output_dir / source_type
            self._shard_writers[source_type] = ShardWriter(output_dir, source_type, self._shard_size)
        return self._shard_writers[source_type]

    async def _emit(
        self,
        source_type: str,
        source_url: str,
        raw_text: str,
        metadata: Mapping[str, Any] | None = None,
        record_id_override: str | None = None,
    ) -> bool:
        """Run the license gate and write a record. Returns True if emitted."""
        ok, license_id = self._license_gate.check(source_url)
        if not ok:
            return False

        record = build_record(
            source_type=source_type,
            source_url=source_url,
            raw_text=raw_text,
            metadata=metadata,
            license_id=license_id,
        )
        if record_id_override is not None:
            record["id"] = record_id_override

        writer = self._get_shard_writer(source_type)
        writer.write(record)
        return True

    async def _process_web(self, source_url: str) -> int:
        """Process a single web URL. Returns 1 if emitted, 0 if dropped."""
        result = await self._web_extractor.extract(source_url)
        if not result.get("raw_text"):
            logger.info("No text extracted from web URL %s — skipping", source_url)
            return 0
        emitted = await self._emit("web", result["source_url"], result["raw_text"], result.get("metadata"))
        return 1 if emitted else 0

    async def _process_document(self, source_ref: str) -> int:
        """Process a single document path. Returns 1 if emitted, 0 if dropped."""
        result = await self._document_extractor.extract(source_ref)
        if not result.get("raw_text"):
            logger.info("No text extracted from document %s — skipping", source_ref)
            return 0
        emitted = await self._emit("document", result["source_url"], result["raw_text"], result.get("metadata"))
        return 1 if emitted else 0

    async def _process_api(self, source_url: str) -> int:
        """Process an API source. Returns count of emitted records."""
        if self._api_extractor is None:
            logger.error("API extractor not configured — cannot process %s", source_url)
            return 0
        count = 0
        async for item in self._api_extractor.extract(source_url):
            if not item.get("raw_text"):
                continue
            record_id = item.pop("id_override", None)
            emitted = await self._emit("api", item["source_url"], item["raw_text"], item.get("metadata"), record_id)
            if emitted:
                count += 1
        return count

    async def _process_youtube(self, source_url: str) -> int:
        """Process a single YouTube URL. Returns 1 if emitted, 0 if dropped."""
        result = await self._youtube_extractor.extract(source_url)
        if not result.get("raw_text"):
            logger.info("No transcript for YouTube video %s — skipping", source_url)
            return 0
        emitted = await self._emit("youtube", result["source_url"], result["raw_text"], result.get("metadata"))
        return 1 if emitted else 0

    async def ingest(self, sources: Sequence[Mapping[str, str]]) -> dict[str, int]:
        """Process a batch of sources and return per-type counts.

        Each source dict must have ``source_type`` and ``source_url`` keys.
        Additional keys are passed as parameters to the extractor.
        """
        counts: dict[str, int] = {st: 0 for st in self.SOURCE_TYPES}
        counts["dropped"] = 0

        for source in sources:
            source_type = source.get("source_type", "")
            source_url = source.get("source_url") or source.get("source_ref") or ""

            if source_type not in self.SOURCE_TYPES:
                logger.warning("Unknown source_type '%s' — skipping %s", source_type, source_url)
                counts["dropped"] += 1
                continue

            if not source_url:
                logger.warning("Empty source_url for source_type=%s — skipping", source_type)
                counts["dropped"] += 1
                continue

            try:
                if source_type == "web":
                    emitted = await self._process_web(source_url)
                elif source_type == "document":
                    emitted = await self._process_document(source_url)
                elif source_type == "api":
                    emitted = await self._process_api(source_url)
                elif source_type == "youtube":
                    emitted = await self._process_youtube(source_url)
                else:
                    emitted = 0

                counts[source_type] += emitted
                if emitted == 0:
                    counts["dropped"] += 1
            except Exception as exc:  # noqa: BLE001
                logger.error("Error processing %s source %s: %s", source_type, source_url, exc)
                counts["dropped"] += 1

        return counts

    async def close(self) -> None:
        """Close all extractors and shard writers."""
        await self._web_extractor.close()
        await self._document_extractor.close()
        if self._api_extractor is not None:
            await self._api_extractor.close()
        await self._youtube_extractor.close()

        for writer in self._shard_writers.values():
            writer.close()

    @property
    def dropped_count(self) -> int:
        return self._license_gate.dropped_count

    @property
    def total_written(self) -> int:
        return sum(w.total_written for w in self._shard_writers.values())


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------


def main() -> None:  # pragma: no cover
    """Run the ingest router from the command line."""
    import argparse

    parser = argparse.ArgumentParser(description="Ingest router — Stage 0 of the curation pipeline")
    parser.add_argument("--manifest", default=DEFAULT_MANIFEST_PATH, help="Path to source manifest YAML")
    parser.add_argument("--output-dir", default=str(DEFAULT_RAW_OUTPUT_DIR), help="Raw output directory")
    parser.add_argument("--shard-size", type=int, default=SHARD_SIZE, help="Records per shard")
    parser.add_argument("--sources-file", help="JSON file with list of source dicts")
    args = parser.parse_args()

    logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(name)s: %(message)s")

    if args.sources_file:
        sources = json.loads(Path(args.sources_file).read_text(encoding="utf-8"))
    else:
        parser.error("--sources-file is required")

    router = IngestRouter(
        manifest_path=args.manifest,
        raw_output_dir=Path(args.output_dir),
        shard_size=args.shard_size,
    )

    async def _run():
        counts = await router.ingest(sources)
        await router.close()
        return counts

    counts = asyncio.run(_run())

    logger.info("Ingest complete: %s", json.dumps(counts))


if __name__ == "__main__":
    main()
