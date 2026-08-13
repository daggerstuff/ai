"""Unit tests for the ingest router — Stage 0 of the curation pipeline.

Tests cover:
- Router dispatch (mock extractors, verify correct one called per source_type).
- Shard writing (verify sharding, file naming, JSONL validity).
- License gate (valid + invalid license paths).
- Web extractor rate limiting (verify per-domain semaphore timing).
- DOCX extraction (fixture .docx with known content).
"""

from __future__ import annotations

import asyncio
import json
import time
from pathlib import Path
from typing import Any
from unittest.mock import AsyncMock, MagicMock, patch
from urllib.robotparser import RobotFileParser

import httpx
import pytest

from training.ingest_router import (
    APIExtractor,
    APIExtractorConfig,
    DocumentExtractor,
    IngestRouter,
    LicenseGate,
    ManifestEntry,
    ShardWriter,
    WebExtractor,
    YouTubeExtractor,
    build_record,
    load_source_manifest,
    lookup_license,
)

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

REPO_ROOT = Path(__file__).resolve().parent.parent.parent


def _allow_all_robots() -> RobotFileParser:
    rp = RobotFileParser()
    rp.parse(["User-agent: *", "Allow: /"])
    return rp


@pytest.fixture
def manifest_path(tmp_path: Path) -> Path:
    """Create a temporary source manifest with known entries."""
    manifest = tmp_path / "source_manifest.yaml"
    manifest.write_text(
        """
sources:
  - url: "https://en.wikipedia.org"
    license: "CC-BY-SA-4.0"
    source_type: "web"
    provenance:
      publisher: "Wikimedia Foundation"
  - url: "https://arxiv.org"
    license: "NOASSERTION"
    source_type: "document"
    provenance:
      publisher: "arXiv"
  - url: "https://api.openalex.org"
    license: "CC0-1.0"
    source_type: "api"
    provenance:
      publisher: "OpenAlex"
  - url: "https://www.youtube.com"
    license: "NOASSERTION"
    source_type: "youtube"
    provenance:
      publisher: "Google LLC"
  - url: "https://unlicensed.example.com"
    license: "GPL-3.0"
    source_type: "web"
    provenance:
      publisher: "Test"
""",
        encoding="utf-8",
    )
    return manifest


@pytest.fixture
def output_dir(tmp_path: Path) -> Path:
    return tmp_path / "raw"


@pytest.fixture
def manifest(manifest_path: Path) -> dict[str, ManifestEntry]:
    return load_source_manifest(manifest_path)


# ---------------------------------------------------------------------------
# Source manifest tests
# ---------------------------------------------------------------------------


class TestSourceManifest:
    def test_load_manifest_finds_entries(self, manifest: dict[str, ManifestEntry]) -> None:
        assert len(manifest) == 5
        assert "https://en.wikipedia.org" in manifest
        assert manifest["https://en.wikipedia.org"].license_id == "CC-BY-SA-4.0"

    def test_load_manifest_missing_file(self, tmp_path: Path) -> None:
        result = load_source_manifest(tmp_path / "nonexistent.yaml")
        assert result == {}

    def test_lookup_license_matches_domain(self, manifest: dict[str, ManifestEntry]) -> None:
        assert lookup_license("https://en.wikipedia.org/wiki/Therapy", manifest) == "CC-BY-SA-4.0"
        assert lookup_license("https://arxiv.org/abs/2401.00001", manifest) == "NOASSERTION"
        assert lookup_license("https://api.openalex.org/works", manifest) == "CC0-1.0"

    def test_lookup_license_falls_back_to_noassertion(self, manifest: dict[str, ManifestEntry]) -> None:
        assert lookup_license("https://unknown.example.com/page", manifest) == "NOASSERTION"

    def test_manifest_entry_matches(self) -> None:
        entry = ManifestEntry(url_or_domain="example.com", license_id="MIT")
        assert entry.matches("https://example.com/page")
        assert entry.matches("https://sub.example.com/page")
        assert not entry.matches("https://other.com/page")

    def test_file_url_matches_by_path_prefix(self) -> None:
        entry = ManifestEntry(url_or_domain="file:///data/books/", license_id="NOASSERTION")
        assert entry.matches("file:///data/books/clinical_textbook.pdf")
        assert not entry.matches("file:///other/books/book.pdf")

    def test_host_match_rejects_similar_domain(self) -> None:
        entry = ManifestEntry(url_or_domain="example.com", license_id="MIT")
        assert not entry.matches("https://example.com.evil.com/page")


# ---------------------------------------------------------------------------
# License gate tests
# ---------------------------------------------------------------------------


class TestLicenseGate:
    def test_valid_license_passes(self, manifest: dict[str, ManifestEntry]) -> None:
        gate = LicenseGate(manifest)
        ok, license_id = gate.check("https://en.wikipedia.org/wiki/Test")
        assert ok is True
        assert license_id == "CC-BY-SA-4.0"

    def test_invalid_license_drops(self, manifest: dict[str, ManifestEntry]) -> None:
        gate = LicenseGate(manifest)
        ok, _license_id = gate.check("https://unlicensed.example.com/page")
        assert ok is False
        assert gate.dropped_count == 1

    def test_unknown_domain_defaults_to_noassertion(self, manifest: dict[str, ManifestEntry]) -> None:
        gate = LicenseGate(manifest)
        ok, license_id = gate.check("https://random.example.com/page")
        assert ok is True
        assert license_id == "NOASSERTION"


# ---------------------------------------------------------------------------
# Shard writer tests
# ---------------------------------------------------------------------------


class TestShardWriter:
    def test_writes_single_shard(self, output_dir: Path) -> None:
        writer = ShardWriter(output_dir, "web", shard_size=100)
        record = {"id": "test1", "source_type": "web", "raw_text": "hello"}
        writer.write(record)
        writer.close()

        shard_path = output_dir / "shard-00001.jsonl"
        assert shard_path.exists()
        lines = shard_path.read_text(encoding="utf-8").strip().splitlines()
        assert len(lines) == 1
        assert json.loads(lines[0]) == record

    def test_shard_rolling(self, output_dir: Path) -> None:
        writer = ShardWriter(output_dir, "web", shard_size=3)
        for i in range(7):
            writer.write({"id": f"r{i}", "source_type": "web", "raw_text": f"text{i}"})
        writer.close()

        # 7 records, shard_size=3 → shards: 3+3+1
        assert writer.total_written == 7
        assert writer.shard_count == 3

        shard1 = (output_dir / "shard-00001.jsonl").read_text(encoding="utf-8").strip().splitlines()
        shard2 = (output_dir / "shard-00002.jsonl").read_text(encoding="utf-8").strip().splitlines()
        shard3 = (output_dir / "shard-00003.jsonl").read_text(encoding="utf-8").strip().splitlines()

        assert len(shard1) == 3
        assert len(shard2) == 3
        assert len(shard3) == 1

    def test_jsonl_validity(self, output_dir: Path) -> None:
        writer = ShardWriter(output_dir, "api", shard_size=100)
        record = {
            "id": "x1",
            "source_type": "api",
            "source_url": "https://api.example.com",
            "raw_text": "sample text",
            "metadata": {"fetch_ts": "2026-01-01T00:00:00+00:00", "lang": "en"},
            "provenance": {"source_url": "https://api.example.com", "license": "MIT"},
        }
        writer.write(record)
        writer.close()

        shard_path = output_dir / "shard-00001.jsonl"
        for line in shard_path.read_text(encoding="utf-8").strip().splitlines():
            parsed = json.loads(line)
            assert "id" in parsed
            assert "source_type" in parsed
            assert "raw_text" in parsed


# ---------------------------------------------------------------------------
# Record builder tests
# ---------------------------------------------------------------------------


class TestBuildRecord:
    def test_record_has_required_fields(self) -> None:
        record = build_record(
            source_type="web",
            source_url="https://example.com/page",
            raw_text="Hello world content here",
            metadata={"fetch_status": 200},
            license_id="MIT",
        )
        assert record["source_type"] == "web"
        assert record["source_url"] == "https://example.com/page"
        assert record["raw_text"] == "Hello world content here"
        assert record["metadata"]["fetch_status"] == 200
        assert record["metadata"]["license"] == "MIT"
        assert record["metadata"]["fetch_ts"].endswith("+00:00")
        assert "provenance" in record
        assert record["provenance"]["license"] == "MIT"

    def test_record_id_is_deterministic(self) -> None:
        r1 = build_record(source_type="web", source_url="https://example.com", raw_text="same text")
        r2 = build_record(source_type="web", source_url="https://example.com", raw_text="same text")
        assert r1["id"] == r2["id"]

    def test_record_id_differs_for_different_content(self) -> None:
        r1 = build_record(source_type="web", source_url="https://example.com", raw_text="text A")
        r2 = build_record(source_type="web", source_url="https://example.com", raw_text="text B")
        assert r1["id"] != r2["id"]

    def test_lang_detection(self) -> None:
        r_en = build_record(
            source_type="web",
            source_url="https://example.com",
            raw_text="The quick brown fox jumps over the lazy dog and the dog is not happy",
        )
        r_de = build_record(
            source_type="web",
            source_url="https://example.com",
            raw_text="Der Hund ist nicht glücklich und das ist ein Problem das wir lösen müssen",
        )
        assert r_en["metadata"]["lang"] == "en"
        assert r_de["metadata"]["lang"] == "de"


# ---------------------------------------------------------------------------
# Router dispatch tests
# ---------------------------------------------------------------------------


class TestRouterDispatch:
    @pytest.mark.asyncio
    async def test_web_source_dispatched_to_web_extractor(self, manifest_path: Path, output_dir: Path) -> None:
        mock_web = MagicMock(spec=WebExtractor)
        mock_web.extract = AsyncMock(
            return_value={
                "source_url": "https://en.wikipedia.org/wiki/Test",
                "raw_text": "Wikipedia article text",
                "metadata": {"fetch_status": 200, "html_title": "Test - Wikipedia"},
            }
        )
        mock_web.close = AsyncMock()

        router = IngestRouter(
            manifest_path=manifest_path,
            raw_output_dir=output_dir,
            web_extractor=mock_web,
        )
        counts = await router.ingest([{"source_type": "web", "source_url": "https://en.wikipedia.org/wiki/Test"}])
        await router.close()

        mock_web.extract.assert_called_once_with("https://en.wikipedia.org/wiki/Test")
        assert counts["web"] == 1

    @pytest.mark.asyncio
    async def test_youtube_source_dispatched_to_youtube_extractor(self, manifest_path: Path, output_dir: Path) -> None:
        mock_yt = MagicMock(spec=YouTubeExtractor)
        mock_yt.extract = AsyncMock(
            return_value={
                "source_url": "https://www.youtube.com/watch?v=abc123",
                "raw_text": "transcript text here",
                "metadata": {"video_id": "abc123", "duration": 300},
            }
        )
        mock_yt.close = AsyncMock()

        router = IngestRouter(
            manifest_path=manifest_path,
            raw_output_dir=output_dir,
            youtube_extractor=mock_yt,
        )
        counts = await router.ingest(
            [{"source_type": "youtube", "source_url": "https://www.youtube.com/watch?v=abc123"}]
        )
        await router.close()

        mock_yt.extract.assert_called_once_with("https://www.youtube.com/watch?v=abc123")
        assert counts["youtube"] == 1

    @pytest.mark.asyncio
    async def test_unknown_source_type_dropped(self, manifest_path: Path, output_dir: Path) -> None:
        router = IngestRouter(
            manifest_path=manifest_path,
            raw_output_dir=output_dir,
        )
        counts = await router.ingest([{"source_type": "unknown", "source_url": "https://example.com"}])
        await router.close()

        assert counts["dropped"] == 1
        assert counts["web"] == 0

    @pytest.mark.asyncio
    async def test_empty_text_dropped(self, manifest_path: Path, output_dir: Path) -> None:
        mock_web = MagicMock(spec=WebExtractor)
        mock_web.extract = AsyncMock(
            return_value={
                "source_url": "https://en.wikipedia.org/wiki/Empty",
                "raw_text": "",
                "metadata": {"fetch_status": 200},
            }
        )
        mock_web.close = AsyncMock()

        router = IngestRouter(
            manifest_path=manifest_path,
            raw_output_dir=output_dir,
            web_extractor=mock_web,
        )
        counts = await router.ingest([{"source_type": "web", "source_url": "https://en.wikipedia.org/wiki/Empty"}])
        await router.close()

        assert counts["web"] == 0
        assert counts["dropped"] == 1

    @pytest.mark.asyncio
    async def test_invalid_license_dropped(self, manifest_path: Path, output_dir: Path) -> None:
        mock_web = MagicMock(spec=WebExtractor)
        mock_web.extract = AsyncMock(
            return_value={
                "source_url": "https://unlicensed.example.com/page",
                "raw_text": "content here",
                "metadata": {"fetch_status": 200},
            }
        )
        mock_web.close = AsyncMock()

        router = IngestRouter(
            manifest_path=manifest_path,
            raw_output_dir=output_dir,
            web_extractor=mock_web,
        )
        counts = await router.ingest([{"source_type": "web", "source_url": "https://unlicensed.example.com/page"}])
        await router.close()

        assert counts["web"] == 0
        assert counts["dropped"] == 1


# ---------------------------------------------------------------------------
# Shard writing integration tests
# ---------------------------------------------------------------------------


class TestShardWritingIntegration:
    @pytest.mark.asyncio
    async def test_shards_written_to_correct_directory(self, manifest_path: Path, output_dir: Path) -> None:
        mock_web = MagicMock(spec=WebExtractor)
        mock_web.extract = AsyncMock(
            return_value={
                "source_url": "https://en.wikipedia.org/wiki/Test",
                "raw_text": "content",
                "metadata": {"fetch_status": 200},
            }
        )
        mock_web.close = AsyncMock()

        router = IngestRouter(
            manifest_path=manifest_path,
            raw_output_dir=output_dir,
            web_extractor=mock_web,
        )
        await router.ingest([{"source_type": "web", "source_url": "https://en.wikipedia.org/wiki/Test"}])
        await router.close()

        web_dir = output_dir / "web"
        assert web_dir.exists()
        shard_files = list(web_dir.glob("shard-*.jsonl"))
        assert len(shard_files) == 1

        records = []
        for line in shard_files[0].read_text(encoding="utf-8").strip().splitlines():
            records.append(json.loads(line))

        assert len(records) == 1
        assert records[0]["source_type"] == "web"
        assert records[0]["raw_text"] == "content"
        assert "provenance" in records[0]


# ---------------------------------------------------------------------------
# Web extractor rate limiting tests
# ---------------------------------------------------------------------------


class TestWebExtractorRateLimit:
    @pytest.mark.asyncio
    async def test_per_domain_rate_limiting(self) -> None:
        """Verify that two requests to the same domain are spaced by rate_limit_seconds."""
        rate_limit = 0.3  # 300ms for fast test
        extractor = WebExtractor(rate_limit_seconds=rate_limit)

        # Mock the client and robots check
        mock_resp = MagicMock()
        mock_resp.status_code = 200
        mock_resp.text = "<html><title>Test</title><body>Hello world and content here</body></html>"

        mock_client = AsyncMock(spec=httpx.AsyncClient)
        mock_client.is_closed = False
        mock_client.get = AsyncMock(return_value=mock_resp)
        mock_client.aclose = AsyncMock()

        extractor._client = mock_client
        extractor._robots_cache["example.com"] = (time.monotonic(), _allow_all_robots())

        start = time.monotonic()
        await extractor.extract("https://example.com/page1")
        await extractor.extract("https://example.com/page2")
        elapsed = time.monotonic() - start

        # Two requests to same domain should take at least rate_limit seconds
        assert elapsed >= rate_limit * 0.9  # allow small timing variance
        assert mock_client.get.call_count == 2

    @pytest.mark.asyncio
    async def test_different_domains_not_rate_limited(self) -> None:
        """Requests to different domains should not block each other."""
        rate_limit = 0.5
        extractor = WebExtractor(rate_limit_seconds=rate_limit)

        mock_resp = MagicMock()
        mock_resp.status_code = 200
        mock_resp.text = "<html><body>Content</body></html>"

        mock_client = AsyncMock(spec=httpx.AsyncClient)
        mock_client.is_closed = False
        mock_client.get = AsyncMock(return_value=mock_resp)
        mock_client.aclose = AsyncMock()

        extractor._client = mock_client
        extractor._robots_cache["a.com"] = (time.monotonic(), _allow_all_robots())
        extractor._robots_cache["b.com"] = (time.monotonic(), _allow_all_robots())

        start = time.monotonic()
        # Concurrent requests to different domains
        await asyncio.gather(
            extractor.extract("https://a.com/page1"),
            extractor.extract("https://b.com/page1"),
        )
        elapsed = time.monotonic() - start

        # Should be fast — different domains don't share rate limit
        assert elapsed < rate_limit


# ---------------------------------------------------------------------------
# Document extractor tests
# ---------------------------------------------------------------------------


class TestWebExtractorRobotsPathLevel:
    @pytest.mark.asyncio
    async def test_path_specific_disallow_blocks_private(self) -> None:
        extractor = WebExtractor()

        robots_text = """User-agent: *
Disallow: /private/
Allow: /
"""
        page_resp = MagicMock()
        page_resp.status_code = 200
        page_resp.text = "<html><body>Public content</body></html>"

        async def fake_get(url: str, **_kwargs: Any) -> MagicMock:
            if url.endswith("/robots.txt"):
                r = MagicMock()
                r.status_code = 200
                r.text = robots_text
                return r
            return page_resp

        mock_client = AsyncMock(spec=httpx.AsyncClient)
        mock_client.is_closed = False
        mock_client.get = AsyncMock(side_effect=fake_get)
        mock_client.aclose = AsyncMock()
        extractor._client = mock_client

        with patch.object(WebExtractor, "_trafilatura_extract", return_value="Public content"):
            result = await extractor.extract("https://example.com/private/secret")
        assert result["raw_text"] == ""
        assert result["metadata"]["fetch_status"] == 403

        with patch.object(WebExtractor, "_trafilatura_extract", return_value="Public content"):
            result = await extractor.extract("https://example.com/public/page")
        assert result["raw_text"] == "Public content"
        assert result["metadata"]["fetch_status"] == 200


class TestDocumentExtractor:
    @pytest.mark.asyncio
    async def test_extract_html(self, tmp_path: Path) -> None:
        html_file = tmp_path / "test.html"
        html_file.write_text(
            "<html><head><title>Test Page</title></head><body><p>Main content here</p></body></html>",
            encoding="utf-8",
        )

        extractor = DocumentExtractor()
        with patch.object(WebExtractor, "_trafilatura_extract", return_value="Main content here"):
            result = await extractor.extract(str(html_file))
        await extractor.close()

        assert result["source_url"] == str(html_file)
        assert result["metadata"]["format"] == "html"
        assert "content" in result["raw_text"].lower()

    @pytest.mark.asyncio
    async def test_extract_docx(self, tmp_path: Path) -> None:
        """Test DOCX extraction with a fixture created via python-docx."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        docx_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_heading("Test Heading", level=1)
        doc.add_paragraph("This is a test paragraph with content.")
        doc.add_heading("Subheading", level=2)
        doc.add_paragraph("Another paragraph here.")
        doc.save(str(docx_path))

        extractor = DocumentExtractor()
        result = await extractor.extract(str(docx_path))
        await extractor.close()

        assert result["source_url"] == str(docx_path)
        assert result["metadata"]["format"] == "docx"
        assert "Test Heading" in result["raw_text"]
        assert "# Test Heading" in result["raw_text"]
        assert "## Subheading" in result["raw_text"]
        assert "test paragraph" in result["raw_text"]

    @pytest.mark.asyncio
    async def test_extract_unsupported_type(self, tmp_path: Path) -> None:
        bad_file = tmp_path / "test.txt"
        bad_file.write_text("content", encoding="utf-8")

        extractor = DocumentExtractor()
        with pytest.raises(ValueError, match="Unsupported document type"):
            await extractor.extract(str(bad_file))
        await extractor.close()


# ---------------------------------------------------------------------------
# API extractor tests
# ---------------------------------------------------------------------------


class TestAPIExtractor:
    @pytest.mark.asyncio
    async def test_pagination_yields_all_items(self) -> None:
        """Verify cursor-based pagination yields items from multiple pages."""
        config = APIExtractorConfig(
            endpoint="https://api.example.com/data",
            text_field="content",
            items_field="results",
            next_cursor_field="next_cursor",
            page_size=2,
            max_pages=5,
        )

        # Mock responses for 3 pages
        req = httpx.Request("GET", "https://api.example.com/data")
        responses = [
            httpx.Response(
                200,
                json={
                    "results": [
                        {"id": "1", "content": "item one"},
                        {"id": "2", "content": "item two"},
                    ],
                    "next_cursor": "page2",
                },
                request=req,
            ),
            httpx.Response(
                200,
                json={
                    "results": [
                        {"id": "3", "content": "item three"},
                        {"id": "4", "content": "item four"},
                    ],
                    "next_cursor": "page3",
                },
                request=req,
            ),
            httpx.Response(
                200,
                json={
                    "results": [{"id": "5", "content": "item five"}],
                    "next_cursor": None,
                },
                request=req,
            ),
        ]

        extractor = APIExtractor(config)
        mock_client = AsyncMock(spec=httpx.AsyncClient)
        mock_client.is_closed = False
        mock_client.get = AsyncMock(side_effect=responses)
        mock_client.aclose = AsyncMock()
        extractor._client = mock_client

        items = []
        async for item in extractor.extract("https://api.example.com"):
            items.append(item)

        assert len(items) == 5
        assert items[0]["raw_text"] == "item one"
        assert items[0]["metadata"]["api_item_id"] == "1"
        assert items[4]["raw_text"] == "item five"

    @pytest.mark.asyncio
    async def test_stops_on_empty_items(self) -> None:
        config = APIExtractorConfig(
            endpoint="https://api.example.com/data",
            items_field="items",
            max_pages=5,
        )

        extractor = APIExtractor(config)
        mock_client = AsyncMock(spec=httpx.AsyncClient)
        mock_client.is_closed = False
        mock_client.get = AsyncMock(
            return_value=httpx.Response(
                200,
                json={"items": [], "next_cursor": None},
                request=httpx.Request("GET", "https://api.example.com/data"),
            )
        )
        mock_client.aclose = AsyncMock()
        extractor._client = mock_client

        items = []
        async for item in extractor.extract("https://api.example.com"):
            items.append(item)

        assert len(items) == 0

    @pytest.mark.asyncio
    async def test_offset_pagination_yields_all_items(self) -> None:
        """Verify offset-based pagination increments offset and yields all items."""
        config = APIExtractorConfig(
            endpoint="https://api.example.com/data",
            text_field="content",
            items_field="results",
            page_size=2,
            max_pages=5,
            offset_mode=True,
        )

        req = httpx.Request("GET", "https://api.example.com/data")
        responses = [
            httpx.Response(
                200,
                json={"results": [{"id": "1", "content": "a"}, {"id": "2", "content": "b"}]},
                request=req,
            ),
            httpx.Response(
                200,
                json={"results": [{"id": "3", "content": "c"}, {"id": "4", "content": "d"}]},
                request=req,
            ),
            httpx.Response(200, json={"results": [{"id": "5", "content": "e"}]}, request=req),
        ]

        extractor = APIExtractor(config)
        mock_client = AsyncMock(spec=httpx.AsyncClient)
        mock_client.is_closed = False
        mock_client.get = AsyncMock(side_effect=responses)
        mock_client.aclose = AsyncMock()
        extractor._client = mock_client

        items = []
        async for item in extractor.extract("https://api.example.com"):
            items.append(item)

        assert len(items) == 5
        offsets_used = [call.kwargs.get("params", {}).get("offset") for call in mock_client.get.call_args_list]
        assert offsets_used == ["0", "2", "4"]

    @pytest.mark.asyncio
    async def test_offset_pagination_stops_on_short_page(self) -> None:
        """Verify offset pagination stops when a page returns fewer items than page_size."""
        config = APIExtractorConfig(
            endpoint="https://api.example.com/data",
            text_field="content",
            items_field="results",
            page_size=2,
            max_pages=5,
            offset_mode=True,
        )

        req = httpx.Request("GET", "https://api.example.com/data")
        responses = [
            httpx.Response(
                200,
                json={"results": [{"id": "1", "content": "a"}, {"id": "2", "content": "b"}]},
                request=req,
            ),
            httpx.Response(200, json={"results": [{"id": "3", "content": "c"}]}, request=req),
        ]

        extractor = APIExtractor(config)
        mock_client = AsyncMock(spec=httpx.AsyncClient)
        mock_client.is_closed = False
        mock_client.get = AsyncMock(side_effect=responses)
        mock_client.aclose = AsyncMock()
        extractor._client = mock_client

        items = []
        async for item in extractor.extract("https://api.example.com"):
            items.append(item)

        assert len(items) == 3
        assert mock_client.get.await_count == 2


# ---------------------------------------------------------------------------
# YouTube extractor tests
# ---------------------------------------------------------------------------


class TestYouTubeExtractor:
    def test_extract_video_id_standard(self) -> None:
        assert YouTubeExtractor._extract_video_id("https://www.youtube.com/watch?v=abc123") == "abc123"

    def test_extract_video_id_short(self) -> None:
        assert YouTubeExtractor._extract_video_id("https://youtu.be/abc123") == "abc123"

    def test_extract_video_id_embed(self) -> None:
        assert YouTubeExtractor._extract_video_id("https://www.youtube.com/embed/abc123") == "abc123"

    def test_fetch_uses_yt_dlp_api_and_unique_temp_directory(self, tmp_path: Path) -> None:
        subtitle = tmp_path / "subtitle.vtt"
        subtitle.write_text(
            "WEBVTT\n\n00:00:00.000 --> 00:00:02.000\nHello world\n",
            encoding="utf-8",
        )

        class FakeYDL:
            def __init__(self, opts: dict[str, Any]):
                self.opts = opts

            def __enter__(self):
                return self

            def __exit__(self, *_args: object):
                return False

            def extract_info(self, _url: str, _download: bool = True) -> dict[str, Any]:
                return {"duration": 123}

        with (
            patch("training.ingest_router.yt_dlp.YoutubeDL", side_effect=FakeYDL),
            patch("training.ingest_router.glob.glob", return_value=[str(subtitle)]),
        ):
            extractor = YouTubeExtractor()
            result = extractor._fetch_transcript_sync("VIDEOID12345", "https://www.youtube.com/watch?v=VIDEOID12345")

        assert result["metadata"]["duration"] == 123
        assert "Hello world" in result["raw_text"]

    def test_invalid_video_id_returns_error(self) -> None:
        extractor = YouTubeExtractor()
        result = extractor._fetch_transcript_sync("bad<id", "https://www.youtube.com/watch?v=bad<id")
        assert result["metadata"]["error"] == "invalid video id"

    def test_invalid_video_id_lengths_rejected(self) -> None:
        extractor = YouTubeExtractor()
        for invalid_id in ("ABCDE12345", "ABCDE1234567"):
            result = extractor._fetch_transcript_sync(invalid_id, f"https://www.youtube.com/watch?v={invalid_id}")
            assert result["metadata"]["error"] == "invalid video id"

    def test_clean_subtitle_strips_formatting(self) -> None:
        vtt_text = """WEBVTT

00:00:00.000 --> 00:00:02.000
Hello world

00:00:02.000 --> 00:00:04.000
This is a test
"""
        result = YouTubeExtractor._clean_subtitle(vtt_text)
        assert "WEBVTT" not in result
        assert "00:00:00.000" not in result
        assert "Hello world" in result
        assert "This is a test" in result
